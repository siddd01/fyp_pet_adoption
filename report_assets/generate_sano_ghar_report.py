from __future__ import annotations

import math
import textwrap
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DIAGRAM_DIR = ROOT / "report_assets" / "diagrams"
OUTPUT_DIR = ROOT / "report_output"
DOCX_PATH = OUTPUT_DIR / "sano_ghar_fyp_report.docx"

THEME = {
    "deep_teal": "1F5C5B",
    "teal": "2F7D79",
    "sand": "E8D7BE",
    "warm": "B97833",
    "ink": "222222",
    "muted": "5B5B5B",
    "light": "F8F5F0",
    "line": "D8CEC0",
    "success": "DFF1E6",
    "warning": "FFF4DB",
}


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D8CEC0", size=6):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.18):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text_runs(paragraph, text: str, font_size=11, color="222222"):
    for idx, chunk in enumerate(text.split("**")):
        run = paragraph.add_run(chunk)
        run.font.name = "Calibri"
        run.font.size = Pt(font_size)
        run.font.color.rgb = rgb(color)
        run.bold = idx % 2 == 1


def add_body_paragraph(doc: Document, text: str, style="Normal", after=8):
    p = doc.add_paragraph(style=style)
    add_text_runs(p, text)
    set_paragraph_spacing(p, after=after)
    return p


def add_bullet_list(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        add_text_runs(p, item)
        set_paragraph_spacing(p, after=3)


def add_numbered_list(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        add_text_runs(p, item)
        set_paragraph_spacing(p, after=3)


def add_heading(doc: Document, level: int, text: str):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    run.font.name = "Cambria"
    run.font.color.rgb = rgb(THEME["deep_teal"])
    if level == 1:
        run.font.size = Pt(17)
        set_paragraph_spacing(p, before=10, after=8)
    elif level == 2:
        run.font.size = Pt(13)
        set_paragraph_spacing(p, before=8, after=6)
    else:
        run.font.size = Pt(11.5)
        set_paragraph_spacing(p, before=6, after=4)
    return p


def apply_document_styles(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(THEME["ink"])

    body = styles["Body Text"] if "Body Text" in styles else styles["Normal"]
    body.font.name = "Calibri"
    body.font.size = Pt(11)

    for level in (1, 2, 3):
        style = styles[f"Heading {level}"]
        style.font.name = "Cambria"
        style.font.color.rgb = rgb(THEME["deep_teal"])
        style.font.bold = True


def cover_page(doc: Document):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(6.6)
    set_table_borders(table, color=THEME["light"], size=0)
    cell = table.cell(0, 0)
    set_cell_shading(cell, THEME["deep_teal"])
    set_cell_margins(cell, top=220, start=260, bottom=220, end=260)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = p.add_run("SANO GHAR\n")
    title.font.name = "Cambria"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = rgb("FFFFFF")
    subtitle = p.add_run("Final Year Project Report\n")
    subtitle.font.name = "Cambria"
    subtitle.font.size = Pt(15)
    subtitle.font.bold = True
    subtitle.font.color.rgb = rgb("F6EBD9")
    small = p.add_run("Web-Based Pet Adoption Center with Integrated Store, Donation, Community and Administrative Management")
    small.font.name = "Calibri"
    small.font.size = Pt(12)
    small.font.color.rgb = rgb("F3F3F3")

    doc.add_paragraph("")
    meta = doc.add_table(rows=6, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.columns[0].width = Inches(2.0)
    meta.columns[1].width = Inches(4.6)
    set_table_borders(meta, color=THEME["line"], size=6)
    labels = [
        ("Student Name", "[Insert Student Name]"),
        ("Student ID", "[Insert Student ID]"),
        ("Programme", "[Insert Programme / Department]"),
        ("Institution", "[Insert College / University]"),
        ("Supervisor", "[Insert Supervisor Name]"),
        ("Submission Date", "04 May 2026"),
    ]
    for row, (label, value) in zip(meta.rows, labels):
        left, right = row.cells
        set_cell_margins(left)
        set_cell_margins(right)
        set_cell_shading(left, THEME["sand"])
        lp = left.paragraphs[0]
        rp = right.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lrun = lp.add_run(label)
        lrun.font.bold = True
        lrun.font.name = "Calibri"
        lrun.font.size = Pt(10.5)
        rrun = rp.add_run(value)
        rrun.font.name = "Calibri"
        rrun.font.size = Pt(10.5)

    doc.add_paragraph("")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text_runs(
        note,
        "Prepared in alignment with the FYP final report template shared by the student. "
        "AI-specific template points are **not applicable** because this version of Sano Ghar relies on rule-based workflows rather than machine learning models.",
        font_size=9.5,
        color=THEME["muted"],
    )
    set_paragraph_spacing(note, after=0)


def declaration_page(doc: Document):
    doc.add_page_break()
    add_heading(doc, 1, "Title and Declaration Sheet")
    add_body_paragraph(
        doc,
        "Project Title: **Sano Ghar: Web-Based Pet Adoption Center with Integrated Store, Donation, Community and Administrative Management**"
    )
    add_body_paragraph(
        doc,
        "I declare that this report describes my own academic work and that all secondary material used in the preparation of this document has been acknowledged in the references and bibliography. "
        "I understand that the final submission should be reviewed by the student and supervisor before institutional submission."
    )
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.5)
    set_table_borders(table, color=THEME["line"], size=6)
    items = [
        ("Student Name", "[Insert Student Name]"),
        ("Signature", "____________________________"),
        ("Supervisor", "[Insert Supervisor Name]"),
        ("Date", "____________________________"),
    ]
    for row, (label, value) in zip(table.rows, items):
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)
        for cell in row.cells:
            set_cell_margins(cell)


def abstract_page(doc: Document):
    doc.add_page_break()
    add_heading(doc, 1, "Abstract")
    add_body_paragraph(
        doc,
        "This report presents **Sano Ghar**, a full-stack web application designed to support the end-to-end operation of a pet adoption center. "
        "The project responds to common operational problems in animal welfare environments: fragmented communication, low visibility of adoptable pets, manual follow-up for applications, disconnected donation records, and weak traceability between user actions and administrative decisions. "
        "The implemented system combines OTP-based account onboarding, pet listing and profile management, adoption applications, role-based staff and administrator workspaces, a product store, Khalti payment integration, direct charitable donation, community engagement posts, issue reporting, notifications, and administrative analytics into a single platform."
    )
    add_body_paragraph(
        doc,
        "The artefact is implemented with React, Vite, Tailwind CSS, Express.js, MySQL, JWT-based authentication, Cloudinary media upload, Nodemailer email delivery, and Khalti payment services. "
        "A Scrum-based development approach was selected because the project contains multiple interacting modules whose requirements are easier to refine across iterations than in a rigid one-pass model. "
        "The report documents the problem domain, research grounding, project methodology, technology decisions, subsystem designs, and implementation evidence derived from the current codebase."
    )
    add_body_paragraph(
        doc,
        "The outcome is an integrated digital platform that improves discoverability of pets, reduces duplicated administrative work, strengthens transaction traceability, and supports both adoption and welfare fundraising within the same system boundary. "
        "Although this version does not include artificial intelligence or mobile-native deployment, it establishes a stable operational foundation that can be extended with future matching, analytics, and multi-channel engagement features."
    )


def contents_page(doc: Document, figure_titles: list[str]):
    doc.add_page_break()
    add_heading(doc, 1, "Contents")
    toc_entries = [
        "1 Introduction",
        "1.1 Project Briefing",
        "1.2 Aim",
        "1.3 Objectives",
        "1.4 Artefact Overview",
        "1.5 Academic Question",
        "1.6 Scope and Limitations",
        "1.7 Report Structure",
        "2 Literature Review",
        "3 Project Methodology",
        "4 Technologies and Tools Used for the Project",
        "5 Artefact Designs",
        "5.1 Subsystem 1: User Onboarding and Pet Adoption",
        "5.2 Subsystem 2: Store, Donation, Community and Reporting",
        "6 Conclusion",
        "7 Critical Evaluation of the Project",
        "8 Evidence of Project Management",
        "9 References and Bibliography",
        "10 Appendices",
    ]
    add_numbered_list(doc, toc_entries)

    add_heading(doc, 2, "Table of Figures")
    add_numbered_list(doc, figure_titles)


@dataclass
class Node:
    id: str
    label: str
    x: int
    y: int
    w: int
    h: int
    kind: str = "round"
    fill: str = "#FFF7ED"
    stroke: str = "#B97833"


@dataclass
class Edge:
    source: str
    target: str
    label: str = ""
    dashed: bool = False


@dataclass
class SequenceMessage:
    src: str
    dst: str
    y: int
    label: str
    dashed: bool = False


@dataclass
class GraphDiagram:
    key: str
    title: str
    width: int
    height: int
    nodes: list[Node] = field(default_factory=list)
    edges: list[Edge] = field(default_factory=list)


@dataclass
class SequenceDiagram:
    key: str
    title: str
    width: int
    height: int
    actors: list[Node]
    messages: list[SequenceMessage]


def create_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        ("C:/Windows/Fonts/cambria.ttc", False),
        ("C:/Windows/Fonts/calibri.ttf", False),
        ("C:/Windows/Fonts/calibrib.ttf", True),
        ("C:/Windows/Fonts/arial.ttf", False),
        ("C:/Windows/Fonts/arialbd.ttf", True),
    ]
    preferred = []
    if bold:
        preferred.extend([("C:/Windows/Fonts/calibrib.ttf", True), ("C:/Windows/Fonts/arialbd.ttf", True)])
    else:
        preferred.extend([("C:/Windows/Fonts/calibri.ttf", False), ("C:/Windows/Fonts/arial.ttf", False)])
    preferred.extend(candidates)
    seen = set()
    for path, _ in preferred:
        if path in seen:
            continue
        seen.add(path)
        try:
            return ImageFont.truetype(path, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


FONT_TITLE = create_font(36, bold=True)
FONT_LABEL = create_font(22, bold=True)
FONT_TEXT = create_font(18)
FONT_SMALL = create_font(16)


def hex_rgb(value: str) -> tuple[int, int, int]:
    value = value.lstrip("#")
    return tuple(int(value[i:i + 2], 16) for i in (0, 2, 4))


def wrap(draw: ImageDraw.ImageDraw, text: str, width: int, font) -> list[str]:
    words = text.replace("\n", " \n ").split()
    lines: list[str] = []
    current = ""
    for word in words:
        if word == "\n":
            lines.append(current.rstrip())
            current = ""
            continue
        trial = f"{current} {word}".strip()
        if draw.textlength(trial, font=font) <= width or not current:
            current = trial
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [text]


def draw_round_box(draw, node: Node):
    fill = hex_rgb(node.fill)
    stroke = hex_rgb(node.stroke)
    draw.rounded_rectangle((node.x, node.y, node.x + node.w, node.y + node.h), radius=22, fill=fill, outline=stroke, width=3)


def draw_ellipse_box(draw, node: Node):
    fill = hex_rgb(node.fill)
    stroke = hex_rgb(node.stroke)
    draw.ellipse((node.x, node.y, node.x + node.w, node.y + node.h), fill=fill, outline=stroke, width=3)


def draw_actor_box(draw, node: Node):
    cx = node.x + node.w // 2
    head_r = 18
    stroke = hex_rgb(node.stroke)
    fill = hex_rgb(node.fill)
    draw.ellipse((cx - head_r, node.y, cx + head_r, node.y + head_r * 2), fill=fill, outline=stroke, width=3)
    body_top = node.y + head_r * 2
    body_bottom = body_top + 40
    draw.line((cx, body_top, cx, body_bottom), fill=stroke, width=3)
    draw.line((cx - 24, body_top + 14, cx + 24, body_top + 14), fill=stroke, width=3)
    draw.line((cx, body_bottom, cx - 24, body_bottom + 32), fill=stroke, width=3)
    draw.line((cx, body_bottom, cx + 24, body_bottom + 32), fill=stroke, width=3)
    label_y = body_bottom + 40
    lines = wrap(draw, node.label, node.w, FONT_SMALL)
    for idx, line in enumerate(lines):
        w = draw.textlength(line, font=FONT_SMALL)
        draw.text((cx - w / 2, label_y + idx * 18), line, fill=stroke, font=FONT_SMALL)


def draw_database_box(draw, node: Node):
    fill = hex_rgb(node.fill)
    stroke = hex_rgb(node.stroke)
    x1, y1, x2, y2 = node.x, node.y, node.x + node.w, node.y + node.h
    ellipse_h = 26
    draw.rectangle((x1, y1 + ellipse_h // 2, x2, y2 - ellipse_h // 2), fill=fill, outline=stroke, width=3)
    draw.ellipse((x1, y1, x2, y1 + ellipse_h), fill=fill, outline=stroke, width=3)
    draw.ellipse((x1, y2 - ellipse_h, x2, y2), outline=stroke, width=3)


def draw_table_box(draw, node: Node):
    fill = hex_rgb("#FFFDFC")
    stroke = hex_rgb(node.stroke)
    header_fill = hex_rgb(node.fill)
    x1, y1, x2, y2 = node.x, node.y, node.x + node.w, node.y + node.h
    draw.rounded_rectangle((x1, y1, x2, y2), radius=16, fill=fill, outline=stroke, width=3)
    header_h = 32
    draw.rounded_rectangle((x1, y1, x2, y1 + header_h), radius=16, fill=header_fill, outline=stroke, width=3)
    lines = node.label.split("\n")
    header = lines[0]
    body_lines = lines[1:]
    draw.text((x1 + 10, y1 + 6), header, fill=hex_rgb("#1F5C5B"), font=FONT_SMALL)
    cur_y = y1 + header_h + 8
    for line in body_lines:
        draw.text((x1 + 10, cur_y), line, fill=hex_rgb("#333333"), font=FONT_SMALL)
        cur_y += 18


def draw_boundary_box(draw, node: Node):
    stroke = hex_rgb(node.stroke)
    draw.rounded_rectangle((node.x, node.y, node.x + node.w, node.y + node.h), radius=20, outline=stroke, width=3)
    draw.text((node.x + 14, node.y + 10), node.label, fill=stroke, font=FONT_SMALL)


def draw_node(draw, node: Node):
    if node.kind == "ellipse":
        draw_ellipse_box(draw, node)
    elif node.kind == "actor":
        draw_actor_box(draw, node)
        return
    elif node.kind == "db":
        draw_database_box(draw, node)
    elif node.kind == "entity":
        draw_table_box(draw, node)
    elif node.kind == "boundary":
        draw_boundary_box(draw, node)
    else:
        draw_round_box(draw, node)

    lines = wrap(draw, node.label, node.w - 22, FONT_SMALL if node.kind == "entity" else FONT_TEXT)
    text_h = len(lines) * 18
    cy = node.y + (node.h - text_h) / 2
    for idx, line in enumerate(lines):
        line_w = draw.textlength(line, font=FONT_TEXT)
        if node.kind == "entity":
            break
        draw.text((node.x + (node.w - line_w) / 2, cy + idx * 20), line, fill=hex_rgb("#2E2E2E"), font=FONT_TEXT)


def arrow(draw, p1, p2, color, dashed=False):
    if dashed:
        dx = p2[0] - p1[0]
        dy = p2[1] - p1[1]
        dist = math.hypot(dx, dy)
        steps = max(1, int(dist // 12))
        for i in range(0, steps, 2):
            s = i / steps
            e = min((i + 1) / steps, 1)
            sx = p1[0] + dx * s
            sy = p1[1] + dy * s
            ex = p1[0] + dx * e
            ey = p1[1] + dy * e
            draw.line((sx, sy, ex, ey), fill=color, width=3)
    else:
        draw.line((p1[0], p1[1], p2[0], p2[1]), fill=color, width=3)
    angle = math.atan2(p2[1] - p1[1], p2[0] - p1[0])
    length = 12
    left = (p2[0] - length * math.cos(angle - math.pi / 6), p2[1] - length * math.sin(angle - math.pi / 6))
    right = (p2[0] - length * math.cos(angle + math.pi / 6), p2[1] - length * math.sin(angle + math.pi / 6))
    draw.polygon([p2, left, right], fill=color)


def edge_points(source: Node, target: Node):
    sx = source.x + source.w / 2
    sy = source.y + source.h / 2
    tx = target.x + target.w / 2
    ty = target.y + target.h / 2
    dx = tx - sx
    dy = ty - sy
    if abs(dx) > abs(dy):
        sx = source.x + source.w if dx > 0 else source.x
        sy = source.y + source.h / 2
        tx = target.x if dx > 0 else target.x + target.w
        ty = target.y + target.h / 2
    else:
        sx = source.x + source.w / 2
        sy = source.y + source.h if dy > 0 else source.y
        tx = target.x + target.w / 2
        ty = target.y if dy > 0 else target.y + target.h
    return (sx, sy), (tx, ty)


def save_graph_png(diagram: GraphDiagram, path: Path):
    image = Image.new("RGB", (diagram.width, diagram.height), color=hex_rgb("#FFFFFF"))
    draw = ImageDraw.Draw(image)
    title_w = draw.textlength(diagram.title, font=FONT_TITLE)
    draw.text(((diagram.width - title_w) / 2, 18), diagram.title, fill=hex_rgb(THEME["deep_teal"]), font=FONT_TITLE)

    by_id = {node.id: node for node in diagram.nodes}
    for node in diagram.nodes:
        if node.kind == "boundary":
            draw_node(draw, node)
    for edge in diagram.edges:
        src = by_id[edge.source]
        dst = by_id[edge.target]
        p1, p2 = edge_points(src, dst)
        arrow(draw, p1, p2, hex_rgb("#7F8C8D"), dashed=edge.dashed)
        if edge.label:
            mx = (p1[0] + p2[0]) / 2
            my = (p1[1] + p2[1]) / 2 - 14
            label_w = draw.textlength(edge.label, font=FONT_SMALL)
            draw.rounded_rectangle((mx - label_w / 2 - 6, my - 4, mx + label_w / 2 + 6, my + 18), radius=8, fill=hex_rgb("#FFFFFF"))
            draw.text((mx - label_w / 2, my), edge.label, fill=hex_rgb(THEME["muted"]), font=FONT_SMALL)
    for node in diagram.nodes:
        if node.kind != "boundary":
            draw_node(draw, node)
    image.save(path)


def save_sequence_png(diagram: SequenceDiagram, path: Path):
    image = Image.new("RGB", (diagram.width, diagram.height), color=hex_rgb("#FFFFFF"))
    draw = ImageDraw.Draw(image)
    title_w = draw.textlength(diagram.title, font=FONT_TITLE)
    draw.text(((diagram.width - title_w) / 2, 18), diagram.title, fill=hex_rgb(THEME["deep_teal"]), font=FONT_TITLE)

    actors = {actor.id: actor for actor in diagram.actors}
    top_y = 90
    bottom_y = diagram.height - 70
    for actor in diagram.actors:
        draw_round_box(draw, actor)
        lines = wrap(draw, actor.label, actor.w - 16, FONT_SMALL)
        y = actor.y + 10
        for line in lines:
            w = draw.textlength(line, font=FONT_SMALL)
            draw.text((actor.x + (actor.w - w) / 2, y), line, fill=hex_rgb("#2E2E2E"), font=FONT_SMALL)
            y += 18
        cx = actor.x + actor.w / 2
        for offset in range(int(top_y), int(bottom_y), 14):
            draw.line((cx, offset, cx, min(offset + 8, bottom_y)), fill=hex_rgb("#8A8A8A"), width=2)

    for message in diagram.messages:
        src = actors[message.src]
        dst = actors[message.dst]
        x1 = src.x + src.w / 2
        x2 = dst.x + dst.w / 2
        y = message.y
        arrow(draw, (x1, y), (x2, y), hex_rgb(THEME["warm"]), dashed=message.dashed)
        label_w = draw.textlength(message.label, font=FONT_SMALL)
        draw.rounded_rectangle((min(x1, x2) + 12, y - 22, min(x1, x2) + 24 + label_w, y - 2), radius=6, fill=hex_rgb("#FFF9F1"))
        draw.text((min(x1, x2) + 18, y - 20), message.label, fill=hex_rgb("#6B4A24"), font=FONT_SMALL)

    image.save(path)


def make_mxfile(name: str, width: int, height: int) -> tuple[ET.Element, ET.Element]:
    mxfile = ET.Element("mxfile", host="app.diagrams.net", modified="2026-05-04T00:00:00.000Z", agent="Codex", version="24.7.17")
    diagram = ET.SubElement(mxfile, "diagram", name=name, id=f"{name}-1")
    model = ET.SubElement(
        diagram,
        "mxGraphModel",
        dx="1200",
        dy="800",
        grid="1",
        gridSize="10",
        guides="1",
        tooltips="1",
        connect="1",
        arrows="1",
        fold="1",
        page="1",
        pageScale="1",
        pageWidth=str(width),
        pageHeight=str(height),
        math="0",
        shadow="0",
    )
    root = ET.SubElement(model, "root")
    ET.SubElement(root, "mxCell", id="0")
    ET.SubElement(root, "mxCell", id="1", parent="0")
    return mxfile, root


def vertex_style(node: Node) -> str:
    base = f"whiteSpace=wrap;html=1;fillColor={node.fill};strokeColor={node.stroke};fontColor=#222222;"
    if node.kind == "ellipse":
        return "ellipse;" + base
    if node.kind == "actor":
        return "shape=umlActor;verticalLabelPosition=bottom;verticalAlign=top;" + base
    if node.kind == "db":
        return "shape=mxgraph.flowchart.database;" + base
    if node.kind == "boundary":
        return "rounded=1;dashed=0;fillColor=none;" + base
    if node.kind == "entity":
        return "rounded=1;align=left;verticalAlign=top;spacing=6;arcSize=12;" + base
    return "rounded=1;arcSize=12;" + base


def save_graph_drawio(diagram: GraphDiagram, path: Path):
    mxfile, root = make_mxfile(diagram.title, diagram.width, diagram.height)
    for node in diagram.nodes:
        cell = ET.SubElement(root, "mxCell", id=node.id, value=node.label, style=vertex_style(node), vertex="1", parent="1")
        ET.SubElement(cell, "mxGeometry", {"x": str(node.x), "y": str(node.y), "width": str(node.w), "height": str(node.h), "as": "geometry"})
    for idx, edge in enumerate(diagram.edges, start=1):
        style = "endArrow=block;html=1;rounded=0;strokeColor=#7F8C8D;"
        if edge.dashed:
            style += "dashed=1;"
        edge_cell = ET.SubElement(root, "mxCell", id=f"e{idx}", value=edge.label, style=style, edge="1", parent="1", source=edge.source, target=edge.target)
        ET.SubElement(edge_cell, "mxGeometry", {"relative": "1", "as": "geometry"})
    ET.ElementTree(mxfile).write(path, encoding="utf-8", xml_declaration=True)


def save_sequence_drawio(diagram: SequenceDiagram, path: Path):
    mxfile, root = make_mxfile(diagram.title, diagram.width, diagram.height)
    for actor in diagram.actors:
        cell = ET.SubElement(root, "mxCell", id=actor.id, value=actor.label, style=vertex_style(actor), vertex="1", parent="1")
        ET.SubElement(cell, "mxGeometry", {"x": str(actor.x), "y": str(actor.y), "width": str(actor.w), "height": str(actor.h), "as": "geometry"})
        lifeline = ET.SubElement(
            root,
            "mxCell",
            id=f"{actor.id}-line",
            value="",
            style="dashed=1;strokeColor=#8A8A8A;endArrow=none;html=1;",
            edge="1",
            parent="1",
            source=actor.id,
            target=actor.id,
        )
        geom = ET.SubElement(lifeline, "mxGeometry", {"relative": "1", "as": "geometry"})
        pts = ET.SubElement(geom, "Array", {"as": "points"})
        ET.SubElement(pts, "mxPoint", x=str(actor.x + actor.w / 2), y="120")
        ET.SubElement(pts, "mxPoint", x=str(actor.x + actor.w / 2), y=str(diagram.height - 70))

    for idx, message in enumerate(diagram.messages, start=1):
        src = next(actor for actor in diagram.actors if actor.id == message.src)
        dst = next(actor for actor in diagram.actors if actor.id == message.dst)
        style = "endArrow=block;html=1;strokeColor=#B97833;"
        if message.dashed:
            style += "dashed=1;"
        edge_cell = ET.SubElement(root, "mxCell", id=f"m{idx}", value=message.label, style=style, edge="1", parent="1", source=src.id, target=dst.id)
        geom = ET.SubElement(edge_cell, "mxGeometry", {"relative": "1", "as": "geometry"})
        pts = ET.SubElement(geom, "Array", {"as": "points"})
        ET.SubElement(pts, "mxPoint", x=str(src.x + src.w / 2), y=str(message.y))
        ET.SubElement(pts, "mxPoint", x=str(dst.x + dst.w / 2), y=str(message.y))

    ET.ElementTree(mxfile).write(path, encoding="utf-8", xml_declaration=True)


def diagrams() -> list[GraphDiagram | SequenceDiagram]:
    return [
        GraphDiagram(
            key="figure_1_fdd",
            title="Functional Decomposition Diagram",
            width=1800,
            height=980,
            nodes=[
                Node("root", "Sano Ghar Pet Adoption Center", 610, 110, 580, 100, fill="#DFF1E6", stroke="#2F7D79"),
                Node("n1", "User and Access Management", 120, 340, 260, 86),
                Node("n2", "Pet Adoption Management", 440, 340, 260, 86),
                Node("n3", "Store and Payment", 760, 340, 260, 86),
                Node("n4", "Charity and Community", 1080, 340, 260, 86),
                Node("n5", "Reports and Administration", 1400, 340, 260, 86),
                Node("n11", "Signup, login, OTP,\nprofile, notifications", 80, 580, 320, 110, fill="#FFF9F1"),
                Node("n21", "Pet catalog, details,\nadoption form, review,\nstatus updates", 410, 580, 320, 110, fill="#FFF9F1"),
                Node("n31", "Products, cart,\ncheckout, Khalti verify,\norder handling", 740, 580, 320, 110, fill="#FFF9F1"),
                Node("n41", "Direct donation,\nimpact posts, likes,\ncomments, transparency", 1070, 580, 320, 110, fill="#FFF9F1"),
                Node("n51", "Issue reports, staff,\nadmin analytics, user\nmanagement", 1400, 580, 320, 110, fill="#FFF9F1"),
            ],
            edges=[
                Edge("root", "n1"),
                Edge("root", "n2"),
                Edge("root", "n3"),
                Edge("root", "n4"),
                Edge("root", "n5"),
                Edge("n1", "n11"),
                Edge("n2", "n21"),
                Edge("n3", "n31"),
                Edge("n4", "n41"),
                Edge("n5", "n51"),
            ],
        ),
        GraphDiagram(
            key="figure_2_architecture",
            title="System Architecture",
            width=1800,
            height=1080,
            nodes=[
                Node("u1", "User", 180, 180, 180, 120, kind="actor", fill="#E9F7EF", stroke="#2F7D79"),
                Node("u2", "Staff", 180, 430, 180, 120, kind="actor", fill="#E9F7EF", stroke="#2F7D79"),
                Node("u3", "Admin", 180, 680, 180, 120, kind="actor", fill="#E9F7EF", stroke="#2F7D79"),
                Node("f", "React + Vite Frontend\n(Tailwind, Axios,\nReact Router, Context API)", 530, 360, 360, 160, fill="#DFF1E6", stroke="#2F7D79"),
                Node("b", "Express.js API Layer\n(auth, adoption, products,\npayments, charity, reports,\nstaff and admin routes)", 1020, 340, 360, 200, fill="#FFF4DB", stroke="#B97833"),
                Node("db", "MySQL Database", 1470, 150, 210, 120, kind="db", fill="#E9F1FF", stroke="#3F6E9A"),
                Node("cl", "Cloudinary\nmedia storage", 1470, 360, 210, 110, kind="round", fill="#F6EEFF", stroke="#7A5B96"),
                Node("kh", "Khalti ePayment", 1470, 550, 210, 110, kind="round", fill="#F3E9FF", stroke="#6B46C1"),
                Node("em", "SMTP / Nodemailer\nOTP and status emails", 1470, 740, 210, 110, kind="round", fill="#FFF6E5", stroke="#9B6A1C"),
            ],
            edges=[
                Edge("u1", "f", "browse, adopt, buy, donate"),
                Edge("u2", "f", "review pets and orders"),
                Edge("u3", "f", "analytics and control"),
                Edge("f", "b", "REST API requests"),
                Edge("b", "db", "data persistence"),
                Edge("b", "cl", "upload images"),
                Edge("b", "kh", "initiate + verify payment"),
                Edge("b", "em", "send OTP / notifications"),
            ],
        ),
        GraphDiagram(
            key="figure_3_use_case",
            title="Use Case Diagram",
            width=1900,
            height=1200,
            nodes=[
                Node("boundary", "Sano Ghar System Boundary", 420, 120, 1140, 920, kind="boundary", fill="#FFFFFF", stroke="#2F7D79"),
                Node("actor_user", "User", 90, 300, 160, 120, kind="actor", fill="#E9F7EF", stroke="#2F7D79"),
                Node("actor_staff", "Staff", 90, 560, 160, 120, kind="actor", fill="#E9F7EF", stroke="#2F7D79"),
                Node("actor_admin", "Admin", 90, 810, 160, 120, kind="actor", fill="#E9F7EF", stroke="#2F7D79"),
                Node("actor_pay", "Khalti", 1650, 390, 160, 120, kind="actor", fill="#F3E9FF", stroke="#6B46C1"),
                Node("actor_mail", "Email Service", 1650, 700, 160, 120, kind="actor", fill="#FFF6E5", stroke="#9B6A1C"),
                Node("uc1", "Register and verify OTP", 540, 210, 250, 70, kind="ellipse", fill="#FFF9F1"),
                Node("uc2", "Browse pet catalog", 880, 210, 220, 70, kind="ellipse", fill="#FFF9F1"),
                Node("uc3", "Submit adoption application", 1180, 210, 270, 70, kind="ellipse", fill="#FFF9F1"),
                Node("uc4", "Shop and manage cart", 620, 390, 240, 70, kind="ellipse", fill="#FFF9F1"),
                Node("uc5", "Donate to charity", 960, 390, 220, 70, kind="ellipse", fill="#FFF9F1"),
                Node("uc6", "View notifications and reports", 1240, 390, 280, 70, kind="ellipse", fill="#FFF9F1"),
                Node("uc7", "Review adoptions", 640, 610, 220, 70, kind="ellipse", fill="#FFF9F1"),
                Node("uc8", "Handle store orders", 940, 610, 220, 70, kind="ellipse", fill="#FFF9F1"),
                Node("uc9", "Manage community posts", 1230, 610, 250, 70, kind="ellipse", fill="#FFF9F1"),
                Node("uc10", "Manage users and staff", 650, 830, 240, 70, kind="ellipse", fill="#FFF9F1"),
                Node("uc11", "Review analytics and transfers", 980, 830, 300, 70, kind="ellipse", fill="#FFF9F1"),
                Node("uc12", "Send OTP and status email", 1280, 830, 260, 70, kind="ellipse", fill="#FFF9F1"),
                Node("uc13", "Process payment verification", 1230, 510, 300, 70, kind="ellipse", fill="#FFF9F1"),
            ],
            edges=[
                Edge("actor_user", "uc1"),
                Edge("actor_user", "uc2"),
                Edge("actor_user", "uc3"),
                Edge("actor_user", "uc4"),
                Edge("actor_user", "uc5"),
                Edge("actor_user", "uc6"),
                Edge("actor_staff", "uc7"),
                Edge("actor_staff", "uc8"),
                Edge("actor_staff", "uc9"),
                Edge("actor_admin", "uc7"),
                Edge("actor_admin", "uc9"),
                Edge("actor_admin", "uc10"),
                Edge("actor_admin", "uc11"),
                Edge("actor_admin", "uc12"),
                Edge("actor_pay", "uc13"),
                Edge("actor_mail", "uc12"),
            ],
        ),
        GraphDiagram(
            key="figure_4_erd",
            title="Simplified Core ERD",
            width=2000,
            height=1220,
            nodes=[
                Node("users", "users\nPK id\nfirst_name\nlast_name\nemail\nrole_id\nis_verified", 90, 180, 230, 150, kind="entity", fill="#DFF1E6", stroke="#2F7D79"),
                Node("pets", "pets\nPK id\nname\nspecies\nbreed\nhealth_status\nimage_url", 430, 150, 230, 150, kind="entity", fill="#FFF4DB", stroke="#B97833"),
                Node("apps", "adoption_applications\nPK id\nFK user_id\nFK pet_id\nstatus\nfull_name\nphone\njob", 780, 130, 280, 190, kind="entity", fill="#FFF9F1", stroke="#B97833"),
                Node("notif", "user_notifications\nPK id\nFK user_id\ntype\nmessage\nrelated_id\nis_read", 1160, 130, 260, 180, kind="entity", fill="#F3E9FF", stroke="#7A5B96"),
                Node("reports", "reports\nPK id\nFK user_id\nsubject\ncategory\npriority\nstatus", 1510, 150, 250, 170, kind="entity", fill="#FFF9F1", stroke="#B97833"),
                Node("products", "products\nPK id\nname\ncategory\nprice\nstock\nimage_url", 170, 610, 240, 170, kind="entity", fill="#DFF1E6", stroke="#2F7D79"),
                Node("orders", "orders\nPK id\nFK user_id\ntotal_amount\ncharity_amount\nstatus\npidx", 560, 580, 240, 180, kind="entity", fill="#FFF4DB", stroke="#B97833"),
                Node("items", "order_items\nPK id\nFK order_id\nFK product_id\nquantity\nprice_at_purchase", 900, 570, 260, 190, kind="entity", fill="#FFF9F1", stroke="#B97833"),
                Node("donations", "donations\nPK id\nFK user_id\namount\nstatus\npidx\nmessage", 1260, 600, 240, 180, kind="entity", fill="#F3E9FF", stroke="#7A5B96"),
                Node("charity", "charity_posts\nPK id\nFK admin_id\ntitle\ncontent\namount_spent\nimage_url", 1580, 590, 260, 190, kind="entity", fill="#FFF9F1", stroke="#B97833"),
            ],
            edges=[
                Edge("users", "apps", "1..*"),
                Edge("pets", "apps", "1..*"),
                Edge("users", "notif", "1..*"),
                Edge("users", "reports", "1..*"),
                Edge("users", "orders", "1..*"),
                Edge("orders", "items", "1..*"),
                Edge("products", "items", "1..*"),
                Edge("users", "donations", "1..*"),
            ],
        ),
        SequenceDiagram(
            key="figure_5_adoption_sequence",
            title="Adoption Approval Sequence",
            width=1850,
            height=980,
            actors=[
                Node("user", "User", 120, 110, 180, 56, fill="#E9F7EF", stroke="#2F7D79"),
                Node("front", "React Frontend", 420, 110, 220, 56, fill="#E9F7EF", stroke="#2F7D79"),
                Node("api", "Express API", 780, 110, 200, 56, fill="#FFF4DB", stroke="#B97833"),
                Node("db", "MySQL DB", 1110, 110, 180, 56, fill="#E9F1FF", stroke="#3F6E9A"),
                Node("staff", "Staff / Admin", 1410, 110, 220, 56, fill="#FFF9F1", stroke="#B97833"),
                Node("mail", "Email Service", 1680, 110, 150, 56, fill="#FFF6E5", stroke="#9B6A1C"),
            ],
            messages=[
                SequenceMessage("user", "front", 250, "open pet and fill form"),
                SequenceMessage("front", "api", 310, "POST /adoptions/apply"),
                SequenceMessage("api", "db", 370, "create application"),
                SequenceMessage("staff", "front", 470, "review pending request"),
                SequenceMessage("front", "api", 530, "PUT /adoptions/:id/status"),
                SequenceMessage("api", "db", 590, "approve one request"),
                SequenceMessage("api", "db", 650, "auto-reject other pending requests"),
                SequenceMessage("api", "db", 710, "insert user notifications"),
                SequenceMessage("api", "mail", 770, "send approval / rejection emails"),
            ],
        ),
        SequenceDiagram(
            key="figure_6_payment_sequence",
            title="Store Checkout and Khalti Verification Sequence",
            width=1850,
            height=980,
            actors=[
                Node("user", "User", 140, 110, 180, 56, fill="#E9F7EF", stroke="#2F7D79"),
                Node("front", "React Checkout", 460, 110, 220, 56, fill="#E9F7EF", stroke="#2F7D79"),
                Node("api", "Express Payment API", 830, 110, 240, 56, fill="#FFF4DB", stroke="#B97833"),
                Node("db", "MySQL DB", 1190, 110, 180, 56, fill="#E9F1FF", stroke="#3F6E9A"),
                Node("pay", "Khalti Gateway", 1490, 110, 210, 56, fill="#F3E9FF", stroke="#6B46C1"),
            ],
            messages=[
                SequenceMessage("user", "front", 250, "enter shipping details"),
                SequenceMessage("front", "api", 320, "POST /payment/checkout"),
                SequenceMessage("api", "db", 390, "create order + reserve stock"),
                SequenceMessage("api", "pay", 470, "initiate payment request"),
                SequenceMessage("pay", "front", 550, "return payment_url"),
                SequenceMessage("front", "pay", 620, "redirect user to Khalti"),
                SequenceMessage("front", "api", 710, "POST /payment/verify or callback"),
                SequenceMessage("api", "pay", 780, "lookup by pidx"),
                SequenceMessage("api", "db", 850, "mark completed or rollback stock"),
            ],
        ),
    ]


FIGURE_TITLES = [
    "Figure 1. Functional Decomposition Diagram of the Sano Ghar platform",
    "Figure 2. High-level system architecture",
    "Figure 3. Use case diagram for user, staff, admin and external services",
    "Figure 4. Simplified core ERD derived from the implemented schema and queries",
    "Figure 5. Adoption approval sequence",
    "Figure 6. Store checkout and Khalti verification sequence",
]


def generate_diagram_assets():
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    generated = {}
    for diagram in diagrams():
        png_path = DIAGRAM_DIR / f"{diagram.key}.png"
        drawio_path = DIAGRAM_DIR / f"{diagram.key}.drawio"
        if isinstance(diagram, GraphDiagram):
            save_graph_png(diagram, png_path)
            save_graph_drawio(diagram, drawio_path)
        else:
            save_sequence_png(diagram, png_path)
            save_sequence_drawio(diagram, drawio_path)
        generated[diagram.key] = {"png": png_path, "drawio": drawio_path, "title": diagram.title}
    return generated


def figure(doc: Document, title: str, image_path: Path, width=6.55):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    set_paragraph_spacing(p, before=4, after=3)
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(title)
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = rgb(THEME["muted"])
    set_paragraph_spacing(caption, after=10)


def table_from_rows(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table, color=THEME["line"], size=6)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        hdr[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(hdr[idx], THEME["deep_teal"])
        set_cell_margins(hdr[idx], top=100, start=100, bottom=100, end=100)
        p = hdr[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = rgb("FFFFFF")
            run.font.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(9.5)
    for row_data in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = value
            row[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(row[idx], top=90, start=100, bottom=90, end=100)
            p = row[idx].paragraphs[0]
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9.5)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    return table


def gantt_table(doc: Document):
    headers = ["Work Package"] + [f"W{i}" for i in range(1, 17)]
    packages = [
        ("Requirements and project framing", range(1, 3)),
        ("Architecture and schema planning", range(3, 5)),
        ("Authentication and user module", range(5, 7)),
        ("Adoption workflow", range(7, 10)),
        ("Store and payment workflow", range(10, 13)),
        ("Donation, community, reporting", range(12, 15)),
        ("Testing, fixes and documentation", range(15, 17)),
    ]
    table = doc.add_table(rows=1 + len(packages), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=THEME["line"], size=6)
    for idx, head in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = head
        set_cell_shading(cell, THEME["deep_teal"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = rgb("FFFFFF")
            run.font.bold = True
            run.font.size = Pt(8.5)
    for row_idx, (label, active) in enumerate(packages, start=1):
        row = table.rows[row_idx].cells
        row[0].text = label
        row[0].paragraphs[0].runs[0].font.bold = True
        for col in range(1, len(headers)):
            cell = row[col]
            if col in active:
                set_cell_shading(cell, THEME["sand"])
                cell.text = " "
            else:
                set_cell_shading(cell, "FFFFFF")
        for cell in row:
            set_cell_margins(cell, top=70, start=70, bottom=70, end=70)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


def introduction_section(doc: Document, assets):
    add_heading(doc, 1, "1 Introduction")
    add_body_paragraph(
        doc,
        "Animal shelters and pet adoption centers frequently struggle with fragmented records, inconsistent follow-up, and limited public visibility of animals that need homes. "
        "ASPCA shelter statistics for 2024 report that **5.8 million dogs and cats** entered shelters and rescues, while **4.2 million** were adopted, showing the continuing operational pressure on animal welfare organisations [1]. "
        "Although these figures are U.S.-based, they illustrate a broader operational reality: adoption centres need systems that reduce process friction, improve transparency, and help staff respond faster."
    )

    add_heading(doc, 2, "1.1 Project Briefing")
    add_body_paragraph(
        doc,
        "Sano Ghar is a web-based pet adoption center platform developed to digitise the main operational flows of a shelter-style service. "
        "The system serves three major user roles: end users who browse pets and adopt, staff members who review operational tasks, and administrators who manage strategic and financial oversight. "
        "Instead of separating adoption, donations, store purchases, community updates, and complaints into disconnected tools, Sano Ghar unifies them inside a single full-stack system."
    )
    add_body_paragraph(
        doc,
        "At implementation level, the system provides OTP-supported account registration, secure login, pet browsing, detailed pet profiles, adoption application submission, profile management, issue reporting, direct donation, store checkout with Khalti, community interaction around charity impact posts, and role-based staff/admin dashboards. "
        "The current version does **not** include AI matching, classification, or recommendation models. The workflow is driven by deterministic business logic, transactional rules, and notification mechanisms."
    )

    add_heading(doc, 2, "1.2 Aim")
    add_body_paragraph(
        doc,
        "The aim of this project is to design and implement an integrated digital platform that improves the accessibility, transparency, and operational efficiency of a pet adoption center while also supporting fundraising, order management, and community communication."
    )

    add_heading(doc, 2, "1.3 Objectives")
    add_bullet_list(
        doc,
        [
            "Provide a secure onboarding flow using email OTP verification and role-aware authentication.",
            "Allow users to browse pets, view relevant profile information, and submit structured adoption applications.",
            "Support staff and administrators in reviewing, approving, rejecting, and tracking adoption activity.",
            "Integrate a product store, cart workflow, and Khalti-based checkout process with traceable order records.",
            "Enable direct charitable donations and track the welfare contribution generated from store purchases.",
            "Support community engagement through impact posts, likes, comments, and notifications.",
            "Provide issue-reporting and administrative monitoring tools that improve accountability and response handling.",
        ],
    )

    add_heading(doc, 2, "1.4 Artefact Overview")
    add_body_paragraph(
        doc,
        "The artefact is modular rather than monolithic. Figure 1 summarises the functional decomposition of the platform. "
        "The highest level of decomposition separates access control, adoption management, commerce, charity and community engagement, and administrative monitoring. "
        "This decomposition is consistent with the implemented route structure in the Express backend and the role-specific page groups in the React frontend."
    )
    figure(doc, FIGURE_TITLES[0], assets["figure_1_fdd"]["png"])

    add_heading(doc, 2, "1.5 Academic Question")
    add_body_paragraph(
        doc,
        "The academic question guiding the project is as follows:"
    )
    quote = doc.add_paragraph()
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = quote.add_run(
        "How effectively can an integrated web-based platform improve process traceability, user experience, and administrative coordination for a pet adoption center compared with fragmented or manual workflows?"
    )
    run.italic = True
    run.font.name = "Cambria"
    run.font.size = Pt(12)
    run.font.color.rgb = rgb(THEME["deep_teal"])
    set_paragraph_spacing(quote, before=4, after=8)

    add_heading(doc, 2, "1.6 Scope and Limitations")
    add_bullet_list(
        doc,
        [
            "Scope: the project covers web-based account management, pet adoption, store checkout, charitable donation, community engagement, reporting, and role-based administration.",
            "Scope: the implementation includes email-based OTP and notification workflows, transactional payment verification, and image upload support.",
            "Limitation: the system depends on configured third-party services such as Khalti, Cloudinary, SMTP, and a live MySQL database.",
            "Limitation: there is no AI-driven pet recommendation, behaviour prediction, or automated suitability scoring in the current release.",
            "Limitation: the project is web-first and not packaged as a native Android or iOS application.",
        ],
    )

    add_heading(doc, 2, "1.7 Report Structure")
    add_body_paragraph(
        doc,
        "The remainder of this report is structured around the institution's FYP template. "
        "Chapter 2 reviews related research and comparable systems, Chapter 3 explains the methodology, Chapter 4 justifies the technology stack, Chapter 5 presents subsystem designs and implementation evidence, Chapter 6 concludes the work, Chapter 7 evaluates the outcome critically, Chapter 8 summarises project management evidence, and the final chapters provide references and appendices."
    )


def literature_review(doc: Document):
    add_heading(doc, 1, "2 Literature Review")
    add_body_paragraph(
        doc,
        "The literature review for Sano Ghar focuses on four themes that directly influence the system design: the operational pressure on adoption centres, the structure of effective pet profiles, the design of comparable pet adoption applications, and the use of payment-enabled digital fundraising."
    )

    add_heading(doc, 2, "2.1 Pet Adoption Operations and Information Needs")
    add_body_paragraph(
        doc,
        "Bradley and Rajendran examined shelter adoption rates through data-driven analysis and highlighted how variables such as age, breed, size, and location influence adoption outcomes and length of stay [3]. "
        "Their work is useful for this project because it shows that operational decisions in shelters benefit from structured, queryable data rather than informal descriptions. "
        "Even though Sano Ghar does not implement predictive modelling, it stores the kinds of descriptive pet information that can later support analytics, prioritisation, or matchmaking enhancements."
    )
    add_body_paragraph(
        doc,
        "Becerra et al. studied user information needs in online pet adoption profiles and found that adopters care strongly about behavioural and physical characteristics when evaluating pets [2]. "
        "This directly supports the Sano Ghar decision to keep pet records richer than a simple name-and-photo listing. "
        "Fields such as species, breed, health status, behaviour, description, and ownership history make the profile more useful and reduce ambiguity during adoption decisions."
    )

    add_heading(doc, 2, "2.2 Comparable Adoption Systems")
    add_body_paragraph(
        doc,
        "The iPET study presents a web-based pet adoption application assessed through ISO 25010 quality criteria and reported positive user reception [4]. "
        "However, the same study recommends future mobile support and integrated communication features, implying that many adoption systems still stop at the listing-and-inquiry stage. "
        "Sano Ghar extends beyond that narrower model by integrating staff workflows, notifications, order management, issue reporting, and financial transparency into one platform."
    )
    add_body_paragraph(
        doc,
        "This comparison suggests an important design gap: prior pet adoption systems often prioritise discoverability of pets but under-support operational continuity after an adoption inquiry begins. "
        "Sano Ghar therefore treats adoption not as an isolated form submission but as part of a broader ecosystem that includes review, notification, approval, rejection, and coordination between users and staff."
    )

    add_heading(doc, 2, "2.3 Digital Donations and Payment-Enabled Welfare Support")
    add_body_paragraph(
        doc,
        "Krisnandy and Nurajijah describe an online donation information system with a payment gateway and argue that electronic fundraising improves accessibility and administrative flow compared with manual donation collection [5]. "
        "That finding is relevant because animal welfare projects often depend on continuous fundraising as much as on adoption outcomes. "
        "Sano Ghar addresses this by supporting both direct donations and a smaller charity amount derived from store orders, improving the linkage between commerce and welfare support."
    )
    add_body_paragraph(
        doc,
        "Khalti's own ePayment documentation requires payment initiation from the merchant's server side and uses a unique `pidx` value for later verification [10]. "
        "This aligns with Sano Ghar's backend-first checkout flow, where the server creates orders, reserves stock, initiates payment, and then verifies transaction status before finalising or rolling back data changes."
    )

    add_heading(doc, 2, "2.4 Security and Technical Foundations")
    add_body_paragraph(
        doc,
        "The system relies on established web technologies rather than experimental infrastructure. React's component model supports modular UI composition [7], Express formalises route-driven API behaviour [8], and MySQL remains a mature multi-user relational database option for transactional workloads [9]. "
        "For security, the platform uses JWT-based session handling in line with RFC 7519 [11], bcrypt hashing through existing Node tooling, and OTP flows that are broadly consistent with OWASP's guidance around multi-factor and password-related protections [12][13]."
    )

    add_heading(doc, 2, "2.5 Research Gap and Contribution")
    add_body_paragraph(
        doc,
        "The main gap identified across the reviewed literature is integration. "
        "Research and comparable applications show strong attention to the adoption portal itself, but they less frequently combine shelter operations, payment-enabled fundraising, order processing, issue reporting, and transparent community communication within a single web application. "
        "The contribution of Sano Ghar is therefore practical and architectural: it demonstrates how these interconnected functions can be implemented cohesively in a role-based full-stack system."
    )


def methodology(doc: Document):
    add_heading(doc, 1, "3 Project Methodology")
    add_body_paragraph(
        doc,
        "A Scrum-oriented methodology was selected for the Sano Ghar project because the system contains multiple interacting subsystems whose detailed requirements become clearer during implementation rather than entirely before it. "
        "The official Scrum Guide remains the current reference released in November 2020 [6]. "
        "This project benefits from Scrum not because it is fashionable, but because the work involves continuous reprioritisation across frontend design, backend routes, schema adjustments, and service integrations."
    )

    add_heading(doc, 2, "3.1 Why Scrum Fits Sano Ghar")
    add_bullet_list(
        doc,
        [
            "The project includes different stakeholder perspectives: adopters, staff, administrators, and donors. Scrum makes it easier to refine features when those perspectives create new requirements.",
            "Adoption, payments, image upload, notifications, and analytics are tightly connected. Iterative delivery reduces the risk of designing all interfaces too early and discovering incompatibilities later.",
            "Third-party dependencies such as Khalti, SMTP, and Cloudinary introduce integration uncertainty. Shorter iterations make failures more visible and easier to isolate.",
            "The platform benefits from incremental demonstrations: an early pet listing and auth module can be validated before more complex order, donation, and reporting features are added.",
        ],
    )

    add_heading(doc, 2, "3.2 Iterative Delivery Structure")
    rows = [
        ["Sprint 1", "Foundation and access", "Project setup, route structure, user auth, OTP verification, core pet listing"],
        ["Sprint 2", "Adoption operations", "Adoption application flow, notifications, approval logic, staff and admin handling"],
        ["Sprint 3", "Commerce and fundraising", "Products, cart, checkout, Khalti flow, direct donation and store contribution"],
        ["Sprint 4", "Administration and refinement", "Community posts, reporting, analytics, security refinements, documentation"],
    ]
    table_from_rows(doc, ["Sprint", "Focus", "Primary Deliverables"], rows, widths=[1.0, 1.4, 4.0])
    doc.add_paragraph("")
    add_body_paragraph(
        doc,
        "The methodology emphasised continuous refinement rather than static specification. "
        "For example, transactional order handling only becomes fully meaningful once product stock, payment verification, and user history are viewed together. "
        "Likewise, adoption status handling required later decisions about automatic rejection of competing pending applications for the same pet. "
        "An iterative method therefore matches the dependency pattern of the actual codebase."
    )


def technologies(doc: Document):
    add_heading(doc, 1, "4 Technologies and Tools Used for the Project")
    add_body_paragraph(
        doc,
        "Technology choices in Sano Ghar were guided by fitness for purpose, learning value, ecosystem maturity, and compatibility with the project's full-stack architecture."
    )
    rows = [
        ["React 19 + Vite 7", "Frontend interface", "Supports modular component-based screens and fast development/build workflow [7]."],
        ["Tailwind CSS 4 + Framer Motion", "Styling and interface polish", "Allows consistent responsive styling and modern user interaction patterns with low overhead."],
        ["Express 5", "Backend HTTP API", "Route-driven structure fits the domain-oriented controllers used by auth, pets, payment, charity and admin flows [8]."],
        ["MySQL + mysql2", "Relational persistence", "Well suited to multi-user transactional data, foreign keys, and order/adoption/report relationships [9]."],
        ["JWT + bcryptjs", "Authentication and password security", "Supports stateless session handling and hashed credentials following common web security practice [11][12]."],
        ["Nodemailer + SMTP", "OTP and workflow email", "Needed for account verification and adoption status communication."],
        ["Cloudinary + multer", "Media upload", "Offloads image storage and simplifies profile, pet and post image handling."],
        ["Khalti ePayment", "Checkout and donation processing", "Provides a familiar payment flow for Nepal-targeted use cases and supports backend initiation / lookup [10]."],
        ["npm, ESLint and Git", "Quality and project support", "Useful for dependency management, linting, and iterative source control across sprints."],
    ]
    table_from_rows(doc, ["Technology", "Role", "Justification"], rows, widths=[1.6, 1.5, 3.35])

    add_heading(doc, 2, "4.1 Environment Summary")
    add_body_paragraph(
        doc,
        "The frontend uses `VITE_API_BASE_URL` to target the backend API, while the backend environment includes database variables, JWT secrets, Khalti keys, Cloudinary credentials, and SMTP settings. "
        "This separation helps keep sensitive credentials away from the client build and is consistent with the architecture shown later in Figure 2."
    )


def artefact_designs(doc: Document, assets):
    add_heading(doc, 1, "5 Artefact Designs")
    add_body_paragraph(
        doc,
        "This chapter groups the artefact into two main subsystem families. "
        "The grouping reflects the actual implementation more accurately than a chapter-per-route approach and keeps the report aligned with the functional structure already shown in Figure 1."
    )

    add_heading(doc, 2, "5.1 Subsystem 1: User Onboarding and Pet Adoption")
    add_heading(doc, 3, "5.1.1 SRS Summary")
    add_bullet_list(
        doc,
        [
            "The user shall be able to register an account and verify it through email OTP before full activation.",
            "The user shall be able to browse available pets and view profile details relevant to adoption decisions.",
            "The user shall be able to submit, edit, and delete adoption applications for eligible pets.",
            "Staff and administrators shall be able to review applications and update status as approved or rejected.",
            "The system shall prevent duplicate successful adoption approval for the same pet and notify affected users.",
            "The system shall preserve an auditable trail through notifications and status-linked messages.",
        ],
    )

    add_heading(doc, 3, "5.1.2 Design and Modelling")
    add_body_paragraph(
        doc,
        "Figure 3 shows the role and service interactions most relevant to the adoption subsystem. "
        "Users initiate onboarding and application activity, while staff and administrators review operational states. "
        "External email support is required because OTP verification and adoption outcome communication are integral to the workflow rather than optional extras."
    )
    figure(doc, FIGURE_TITLES[2], assets["figure_3_use_case"]["png"])

    add_body_paragraph(
        doc,
        "Figure 4 provides a simplified ERD derived from the implemented schema and SQL queries. "
        "The core adoption path depends on relationships between `users`, `pets`, `adoption_applications`, and `user_notifications`, while related service tables such as `reports` and `charity_posts` extend the user's engagement lifecycle."
    )
    figure(doc, FIGURE_TITLES[3], assets["figure_4_erd"]["png"])

    add_body_paragraph(
        doc,
        "Figure 5 models the approval sequence. A particularly important design decision is the use of transactional logic when an application is approved. "
        "The backend updates the chosen record, automatically rejects competing pending requests for the same pet, inserts user notifications, and then dispatches status emails. "
        "This keeps the business rule consistent and avoids partial approval states."
    )
    figure(doc, FIGURE_TITLES[4], assets["figure_5_adoption_sequence"]["png"])

    add_heading(doc, 3, "5.1.3 Testing and Verification Matrix")
    rows = [
        ["A1", "Signup with OTP verification", "Account is created, OTP is sent, and verification activates the user", "authController.signup + verifyOTP"],
        ["A2", "Submit adoption form", "A structured application is created for the selected pet", "AdoptionForm.jsx + createAdoptionApplication"],
        ["A3", "Prevent duplicate request for same pet", "Existing application is detected and repeat submission is blocked", "AdoptionForm.jsx existingApplication check + unique user/pet constraint"],
        ["A4", "Approve one application for a pet", "Selected application becomes approved and other pending ones are rejected", "updateAdoptionStatus transaction logic"],
        ["A5", "Delete pet with pending applications", "Deletion is blocked until pending requests are resolved", "petController.deletePet / staffPetController"],
    ]
    table_from_rows(doc, ["ID", "Scenario", "Expected Behaviour", "Implementation Evidence"], rows, widths=[0.45, 1.8, 2.55, 1.7])

    add_heading(doc, 2, "5.2 Subsystem 2: Store, Donation, Community and Reporting")
    add_heading(doc, 3, "5.2.1 SRS Summary")
    add_bullet_list(
        doc,
        [
            "Users shall be able to browse products, manage a cart, provide shipping details, and initiate checkout.",
            "The backend shall create order records and reserve stock before handing control to the payment provider.",
            "The system shall verify Khalti transactions and either finalise the order or restore stock on failure.",
            "Users shall be able to donate directly and see community-facing charity impact posts.",
            "Administrators shall be able to review analytics, track donation inflow, manage posts, and manage users.",
            "Users shall be able to submit issue reports and receive notifications when report status changes.",
        ],
    )

    add_heading(doc, 3, "5.2.2 Design and Modelling")
    add_body_paragraph(
        doc,
        "The subsystem depends on several external integrations, so a high-level architecture view is useful. "
        "Figure 2 shows the relationship between the React client, Express backend, MySQL persistence layer, and third-party services for media, payments, and email. "
        "This architecture keeps sensitive actions such as payment initiation and verification on the server side, which is also the flow required by Khalti's documentation [10]."
    )
    figure(doc, FIGURE_TITLES[1], assets["figure_2_architecture"]["png"])

    add_body_paragraph(
        doc,
        "Figure 6 illustrates the checkout and payment lifecycle. "
        "The implementation deliberately creates the order and reserves stock before redirecting to Khalti. "
        "After payment lookup, successful transactions clear the cart and mark the order as completed, while failed or refunded outcomes restore reserved inventory. "
        "This reduces inconsistency between stock levels and payment outcomes."
    )
    figure(doc, FIGURE_TITLES[5], assets["figure_6_payment_sequence"]["png"])

    add_heading(doc, 3, "5.2.3 Testing and Verification Matrix")
    rows = [
        ["C1", "Add products and cart management", "Products can be loaded and cart items are tied to the authenticated user", "productController + cart flow described in FIXES_SUMMARY"],
        ["C2", "Checkout initiation", "Order is created, stock is reserved, and Khalti payment_url is returned", "paymentController.handleCheckout"],
        ["C3", "Payment failure or cancellation", "Pending order is marked failed and reserved stock is restored", "verifyPaymentAndUpdateOrder failure branch"],
        ["C4", "Direct donation verification", "Donation record status updates after Khalti lookup", "charityController.initiateDonation + verifyDonation"],
        ["C5", "Report status notification", "User receives a notification when admin changes report status", "reportController.updateReportStatus"],
        ["C6", "Staff order acceptance", "Paid orders can be accepted and assigned to staff with delivery metadata", "staffOrderController.acceptStaffOrder"],
    ]
    table_from_rows(doc, ["ID", "Scenario", "Expected Behaviour", "Implementation Evidence"], rows, widths=[0.45, 1.8, 2.55, 1.7])


def conclusion(doc: Document):
    add_heading(doc, 1, "6 Conclusion")
    add_body_paragraph(
        doc,
        "Sano Ghar demonstrates that an integrated web platform can improve operational coordination for a pet adoption center by connecting adoption, commerce, fundraising, reporting, and administrative oversight within the same system boundary. "
        "The project addressed the academic question by showing that process traceability becomes stronger when applications, notifications, payments, and role-specific dashboards all share a coherent backend and relational data model."
    )
    add_body_paragraph(
        doc,
        "The implemented system goes beyond a basic adoption portal. It supports structured pet data, protected account onboarding, transactional payment flows, direct donations, community transparency, and issue escalation. "
        "The result is not merely a user-facing website, but an operational platform that reduces fragmentation across day-to-day shelter activities."
    )


def critical_evaluation(doc: Document):
    add_heading(doc, 1, "7 Critical Evaluation of the Project")
    add_heading(doc, 2, "7.1 Findings and Process")
    add_body_paragraph(
        doc,
        "One of the strongest aspects of the project is its breadth of integration. "
        "Many student systems stop after authentication and CRUD, whereas Sano Ghar connects multiple real-world concerns: adoption review, stock-sensitive checkout, donation handling, email notifications, and analytics. "
        "This breadth increased implementation complexity, but it also made the final artefact more realistic and academically richer."
    )

    add_heading(doc, 2, "7.2 System Evaluation")
    add_bullet_list(
        doc,
        [
            "Strength: transactional logic in adoption approval and payment verification reduces inconsistent system states.",
            "Strength: role separation between user, staff, and admin improves maintainability and access clarity.",
            "Strength: the platform supports both mission activity (adoption) and sustainability activity (donation/store contribution).",
            "Weakness: the current system depends on third-party services and careful environment configuration for full end-to-end execution.",
            "Weakness: mobile-native delivery, advanced search, and recommendation support are not yet present.",
            "Weakness: a fuller automated test suite would strengthen long-term maintainability beyond manual or code-level verification.",
        ],
    )

    add_heading(doc, 2, "7.3 Self Reflection")
    add_body_paragraph(
        doc,
        "From a learning perspective, the project provides strong evidence of full-stack growth. "
        "It requires thinking across frontend interaction design, backend route composition, relational data modelling, third-party integration, and user communication flows. "
        "The work also reinforces a key professional lesson: high-value software is usually not about a single feature, but about how several features remain consistent when they interact."
    )


def project_management(doc: Document):
    add_heading(doc, 1, "8 Evidence of Project Management")
    add_heading(doc, 2, "8.1 Milestone Log Summary")
    add_body_paragraph(
        doc,
        "The official signed supervisor log sheet should be attached by the student during final institutional submission. "
        "For report completeness, Table 8.1 summarises the major delivery milestones reflected in the implemented codebase."
    )
    rows = [
        ["1", "Project framing and problem definition", "Adoption center workflow selected; scope expanded to store, donation and reporting support"],
        ["2", "System foundation", "React and Express project structure established; environment variables and API base URL configured"],
        ["3", "User auth and pet modules", "Signup, login, OTP verification, profile support and pet browsing/apply flows added"],
        ["4", "Role-based operations", "Staff and admin dashboards, adoption handling, user management and staff management added"],
        ["5", "Commerce and fundraising", "Products, cart, checkout, Khalti verification, donation flow and charity analytics added"],
        ["6", "Community and reporting", "Impact posts, likes/comments, notifications, reporting workflow and final refinements added"],
    ]
    table_from_rows(doc, ["Milestone", "Focus", "Outcome"], rows, widths=[0.8, 1.8, 4.35])
    doc.add_paragraph("")

    add_heading(doc, 2, "8.2 High-Level Gantt")
    gantt_table(doc)


def references(doc: Document):
    add_heading(doc, 1, "9 References and Bibliography")
    refs = [
        "[1] ASPCA, \"U.S. Animal Shelter Statistics,\" Shelter Intake and Surrender, 2024 statistics page. Accessed: 04 May 2026. Available: https://www.aspca.org/helping-people-pets/shelter-intake-and-surrender",
        "[2] Z. M. Becerra, S. Parmar, K. May, and R. E. Stuck, \"Exploring User Information Needs in Online Pet Adoption Profiles,\" Proceedings of the Human Factors and Ergonomics Society Annual Meeting, vol. 64, no. 1, pp. 1308-1312, 2021. DOI: https://doi.org/10.1177/1071181320641311",
        "[3] J. Bradley and S. Rajendran, \"Increasing adoption rates at animal shelters: a two-phase approach to predict length of stay and optimal shelter allocation,\" BMC Veterinary Research, vol. 17, article 70, 2021. DOI: https://doi.org/10.1186/s12917-020-02728-2",
        "[4] J. E. Gatmaitan, D. V. P. Azurin, J. P. T. Saclolo, J. A. G. Yango, K. P. Lorenzo, C. N. P. Olipas, and R. T. Alegado, \"iPET: The Design, Development, and Assessment of a Web-Based Application for Pet Adoption,\" Formosa Journal of Computer and Information Science, vol. 3, no. 1, 2024. DOI: https://doi.org/10.55927/fjcis.v3i1.6031",
        "[5] A. Krisnandy and N. Nurajijah, \"Online Donation Information System with Payment Gateway,\" International Journal of Information System and Technology, vol. 6, no. 1, 2022. DOI: https://doi.org/10.30645/ijistech.v6i1.204",
        "[6] K. Schwaber and J. Sutherland, \"The Scrum Guide,\" official current version, Nov. 2020. Available: https://scrumguides.org/download.html",
        "[7] React Documentation, \"Quick Start,\" React.dev. Accessed: 04 May 2026. Available: https://react.dev/learn",
        "[8] Express.js Documentation, \"Routing,\" Expressjs.com. Accessed: 04 May 2026. Available: https://expressjs.com/en/guide/routing.html",
        "[9] Oracle, \"MySQL 8.0 Reference Manual,\" Accessed: 04 May 2026. Available: https://dev.mysql.com/doc/refman/8.0/en/introduction.html",
        "[10] Khalti, \"Web Checkout - Khalti Payment Gateway,\" Accessed: 04 May 2026. Available: https://docs.khalti.com/khalti-epayment/",
        "[11] M. Jones, J. Bradley, and N. Sakimura, \"RFC 7519: JSON Web Token (JWT),\" Internet Engineering Task Force, May 2015. Available: https://datatracker.ietf.org/doc/html/rfc7519",
        "[12] OWASP Cheat Sheet Series, \"Password Storage Cheat Sheet,\" Accessed: 04 May 2026. Available: https://cheatsheetseries.owasp.org/cheatsheets/Password_Storage_Cheat_Sheet.html",
        "[13] OWASP Cheat Sheet Series, \"Multifactor Authentication Cheat Sheet,\" Accessed: 04 May 2026. Available: https://cheatsheetseries.owasp.org/cheatsheets/Multifactor_Authentication_Cheat_Sheet.html",
    ]
    for ref in refs:
        add_body_paragraph(doc, ref, after=4)


def appendices(doc: Document):
    add_heading(doc, 1, "10 Appendices")
    add_heading(doc, 2, "10.1 End User Manual")
    add_numbered_list(
        doc,
        [
            "Create an account using signup, then verify the OTP sent to the registered email address.",
            "Log in and browse the available pets from the adoption area.",
            "Open a pet profile, review its details, and submit an adoption application if eligible.",
            "Use the shop to add products to the cart and proceed to checkout with delivery details.",
            "Complete payments or donations through the Khalti redirection flow and return to the application for status verification.",
            "Check notifications to track adoption updates, report updates, and other system messages.",
        ],
    )

    add_heading(doc, 2, "10.2 System Configuration Details")
    rows = [
        ["Frontend", "VITE_API_BASE_URL", "http://localhost:5000/api"],
        ["Backend", "DB_HOST / DB_USER / DB_PASSWORD / DB_NAME", "MySQL connection settings"],
        ["Backend", "JWT_SECRET / ADMIN_JWT_SECRET / STAFF_JWT_SECRET", "Authentication secrets"],
        ["Backend", "KHALTI_SECRET_KEY / KHALTI_TEST_SECRET_KEY / KHALTI_BASE_URL", "Payment gateway credentials"],
        ["Backend", "FRONTEND_URL", "Frontend callback target"],
        ["Backend", "CLOUDINARY_CLOUD_NAME / API_KEY / API_SECRET", "Media upload support"],
        ["Backend", "EMAIL_HOST / EMAIL_PORT / EMAIL_USER / EMAIL_PASS", "OTP and workflow email delivery"],
    ]
    table_from_rows(doc, ["Layer", "Variable(s)", "Purpose"], rows, widths=[0.9, 3.1, 2.95])
    doc.add_paragraph("")
    add_heading(doc, 2, "10.3 Diagram Source Files")
    add_body_paragraph(
        doc,
        "Editable diagram sources for this report were generated in draw.io-compatible format and stored in the workspace under `report_assets/diagrams`. "
        "These files can be opened in diagrams.net / draw.io for further refinement before final academic submission."
    )


def build_docx(assets):
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    apply_document_styles(doc)

    cover_page(doc)
    declaration_page(doc)
    abstract_page(doc)
    contents_page(doc, FIGURE_TITLES)
    introduction_section(doc, assets)
    literature_review(doc)
    methodology(doc)
    technologies(doc)
    artefact_designs(doc, assets)
    conclusion(doc)
    critical_evaluation(doc)
    project_management(doc)
    references(doc)
    appendices(doc)

    doc.save(DOCX_PATH)


def main():
    assets = generate_diagram_assets()
    build_docx(assets)
    print(f"Generated report: {DOCX_PATH}")
    for key, value in assets.items():
        print(f"{key}: {value['drawio']} | {value['png']}")


if __name__ == "__main__":
    main()
